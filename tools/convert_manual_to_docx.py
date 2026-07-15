from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_BREAK
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "SwiftAI_ERP_User_Manual.md"
OUTPUT = ROOT / "docs" / "SwiftAI_ERP_User_Manual.docx"


def clean_inline(text: str) -> str:
    text = re.sub(r"!\[([^\]]*)\]\([^)]+\)", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    text = text.replace("**", "")
    text = text.replace("`", "")
    text = text.replace("&gt;", ">")
    return text.strip()


def add_inline(paragraph, text: str) -> None:
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
        else:
            paragraph.add_run(re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", part))


def is_table_block(lines: list[str], index: int) -> bool:
    return (
        index + 1 < len(lines)
        and lines[index].strip().startswith("|")
        and lines[index + 1].strip().startswith("|")
        and set(lines[index + 1].replace("|", "").strip()) <= {"-", ":", " "}
    )


def parse_table(lines: list[str], index: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    while index < len(lines) and lines[index].strip().startswith("|"):
        line = lines[index].strip()
        cells = [clean_inline(c) for c in line.strip("|").split("|")]
        if not all(set(c) <= {"-", ":", " "} for c in cells):
            rows.append(cells)
        index += 1
    return rows, index


def set_document_style(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10)

    for name, size in [
        ("Heading 1", 18),
        ("Heading 2", 15),
        ("Heading 3", 12),
    ]:
        style = document.styles[name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)


def add_picture(document: Document, alt: str, rel_path: str) -> None:
    image_path = (SOURCE.parent / rel_path).resolve()
    if not image_path.exists():
        document.add_paragraph(f"[Missing screenshot: {alt}]")
        return

    document.add_paragraph(clean_inline(alt), style=None)
    try:
        document.add_picture(str(image_path), width=Inches(6.8))
    except Exception as exc:  # python-docx can embed supported images directly.
        document.add_paragraph(f"[Could not embed image {image_path.name}: {exc}]")


def convert() -> None:
    document = Document()
    set_document_style(document)

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    i = 0
    in_code = False

    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()

        if line.startswith("```"):
            in_code = not in_code
            i += 1
            continue

        if in_code:
            if line.strip():
                p = document.add_paragraph()
                run = p.add_run(line)
                run.font.name = "Consolas"
            i += 1
            continue

        if not line.strip():
            i += 1
            continue

        image = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", line.strip())
        if image:
            add_picture(document, image.group(1), image.group(2))
            i += 1
            continue

        if is_table_block(lines, i):
            rows, i = parse_table(lines, i)
            if rows:
                table = document.add_table(rows=len(rows), cols=max(len(r) for r in rows))
                table.style = "Table Grid"
                for r_idx, row in enumerate(rows):
                    for c_idx, cell in enumerate(row):
                        table.cell(r_idx, c_idx).text = cell
                        if r_idx == 0:
                            for paragraph in table.cell(r_idx, c_idx).paragraphs:
                                for run in paragraph.runs:
                                    run.bold = True
                document.add_paragraph()
            continue

        heading = re.match(r"^(#{1,4})\s+(.+)$", line)
        if heading:
            level = min(len(heading.group(1)), 3)
            text = clean_inline(heading.group(2))
            document.add_heading(text, level=level)
            i += 1
            continue

        if line.startswith(">"):
            p = document.add_paragraph(style=None)
            add_inline(p, clean_inline(line.lstrip("> ")))
            i += 1
            continue

        ordered = re.match(r"^\d+\.\s+(.+)$", line)
        if ordered:
            p = document.add_paragraph(style="List Number")
            add_inline(p, ordered.group(1))
            i += 1
            continue

        if line.lstrip().startswith("- "):
            p = document.add_paragraph(style="List Bullet")
            add_inline(p, line.lstrip()[2:])
            i += 1
            continue

        p = document.add_paragraph()
        add_inline(p, line)
        i += 1

    document.save(OUTPUT)


if __name__ == "__main__":
    convert()
    print(OUTPUT)
